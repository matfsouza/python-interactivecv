from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture(
    'picture.jpg', 
    width=Inches(2.0)
)

# name phone number and email details
name = input('Qual seu nome? ')
phone_number = input('Qual seu número de telefone? ')
email = input('Qual seu email? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

#about me
document.add_heading('Sobre mim')
about_me = input('Me fale sobre você: ')
document.add_paragraph(about_me)

#work experience
document.add_heading('Experiência de Trabalho')
p = document.add_paragraph()

company = input('Coloque sua empresa: ')
from_date =  input('Coloque a data que começou na empresa: ')
to_date = input('Coloque a data que saiu da empresa: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Descreva sua experiencia trabalhando na ' + company + ' '
)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experience = input(
        'Você tem mais alguma experiência? Sim ou Não ')
    if has_more_experience.lower() == 'sim':
        p = document.add_paragraph()

        company = input('Coloque sua empresa ')
        from_date =  input('Coloque quando começou na empresa ')
        to_date = input('Coloque quando saiu dela ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Descreva sua experiencia trabalhando na ' + company + ' '
        )
        p.add_run(experience_details)

    else:
        break

# skills
document.add_heading('Skills')
skills = input('Qual linguagem você sabe e tem experiência? ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

# more skills
while True:
    has_more_skills = input(
        'Você sabe mais alguma linguagem? Sim ou Não ')
    if has_more_skills.lower() == 'sim':
        skills = input('Qual linguagem você sabe e tem experiência? ')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break
        
# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Código feito por Matheus Ferreira'

document.save('cv.docx')